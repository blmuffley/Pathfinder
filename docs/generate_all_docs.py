#!/usr/bin/env python3
"""Generate DOCX versions of all Pathfinder markdown documentation.

Uses python-docx (pure Python, no native deps).
Run: python3 docs/generate_all_docs.py
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(DOCS_DIR)

# All markdown files to convert, grouped by output subfolder
MD_FILES = {
    "guides": [
        "guides/installation-guide.md",
        "guides/operations-runbook.md",
        "guides/implementation-playbook.md",
        "guides/partner-enablement-guide.md",
        "guides/user-guide-workspace.md",
        "guides/five-year-business-case.md",
    ],
    "reference": [
        "reference/api-reference.md",
        "reference/sequence-diagrams.md",
        "reference/security-architecture.md",
        "reference/data-dictionary.md",
        "reference/network-firewall-diagram.md",
        "reference/capacity-planning.md",
        "reference/architecture-decision-records.md",
    ],
    "methodology": [
        "methodology/pathfinder-deployment-methodology.md",
    ],
    ".": [
        # CHANGELOG is at project root
    ],
}

BRAND_GREEN = RGBColor(0x39, 0xFF, 0x14)
BRAND_DARK = RGBColor(0x1C, 0x19, 0x17)
GRAY = RGBColor(0x78, 0x71, 0x6C)


def set_style(doc):
    """Configure Avennorth document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15


def add_header_footer(doc, title):
    """Add Avennorth branding header/footer."""
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "AVENNORTH  |  CONFIDENTIAL"
    hp.style.font.size = Pt(8)
    hp.style.font.color.rgb = GRAY
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Avennorth Pathfinder + Intelligence Platform"
    fp.style.font.size = Pt(8)
    fp.style.font.color.rgb = GRAY
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a styled DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    set_style(doc)

    # Extract title from first # heading
    title = os.path.basename(md_path).replace('.md', '')
    for line in lines:
        if line.startswith('# '):
            title = line[2:].strip()
            break

    add_header_footer(doc, title)

    in_code_block = False
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Parse table
        headers = [c.strip() for c in table_rows[0].split('|')[1:-1]]
        data = []
        for row in table_rows[2:]:  # Skip header + separator
            cells = [c.strip() for c in row.split('|')[1:-1]]
            if cells:
                data.append(cells)

        if headers:
            ncols = len(headers)
            tbl = doc.add_table(rows=1 + len(data), cols=ncols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Header row
            for i, h in enumerate(headers):
                cell = tbl.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(8)

            # Data rows
            for ri, row_data in enumerate(data):
                for ci in range(min(ncols, len(row_data))):
                    cell = tbl.rows[ri + 1].cells[ci]
                    cell.text = row_data[ci]
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)

        in_table = False
        table_rows = []

    for line in lines:
        stripped = line.rstrip('\n')

        # Code blocks
        if stripped.startswith('```'):
            if in_table:
                flush_table()
            in_code_block = not in_code_block
            if in_code_block:
                continue
            else:
                continue

        if in_code_block:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(stripped)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        # Table detection
        if '|' in stripped and stripped.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            continue
        elif in_table:
            flush_table()

        # Headings
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:].strip(), level=4)
            continue

        # Horizontal rule
        if stripped.startswith('---'):
            doc.add_paragraph('_' * 60)
            continue

        # Bullet points
        if stripped.startswith('- [ ] ') or stripped.startswith('- [x] '):
            check = '[ ]' if '[ ]' in stripped else '[x]'
            text = stripped.split('] ', 1)[1] if '] ' in stripped else stripped[6:]
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(check + ' ' + text).font.size = Pt(9)
            continue

        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:].strip()
            # Handle bold
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(9)
                if i % 2 == 1:
                    run.bold = True
            continue

        if re.match(r'^\d+\. ', stripped):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s*', '', stripped)
            p.add_run(text).font.size = Pt(9)
            continue

        # Empty line
        if not stripped.strip():
            doc.add_paragraph('')
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        # Handle inline bold and backticks
        text = stripped.strip()
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x80)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10)

    # Flush any remaining table
    if in_table:
        flush_table()

    doc.save(docx_path)
    return title


def main():
    generated = []

    for subfolder, files in MD_FILES.items():
        for md_file in files:
            md_path = os.path.join(DOCS_DIR, md_file)
            if not os.path.exists(md_path):
                print(f"  SKIP (not found): {md_file}")
                continue

            docx_file = md_file.replace('.md', '.docx')
            docx_path = os.path.join(DOCS_DIR, docx_file)
            os.makedirs(os.path.dirname(docx_path), exist_ok=True)

            title = parse_md_to_docx(md_path, docx_path)
            size_kb = os.path.getsize(docx_path) // 1024
            print(f"  OK: {docx_file} ({size_kb} KB)")
            generated.append(docx_file)

    # Also convert CHANGELOG from project root
    changelog_md = os.path.join(PROJECT_ROOT, 'CHANGELOG.md')
    if os.path.exists(changelog_md):
        changelog_docx = os.path.join(DOCS_DIR, 'CHANGELOG.docx')
        parse_md_to_docx(changelog_md, changelog_docx)
        print(f"  OK: CHANGELOG.docx")
        generated.append('CHANGELOG.docx')

    print(f"\nGenerated {len(generated)} DOCX files.")


if __name__ == '__main__':
    print("Generating DOCX files for Avennorth Pathfinder documentation...\n")
    main()
